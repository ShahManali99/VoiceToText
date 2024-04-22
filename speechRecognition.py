#Create app using with separating widget  
from tkinter import *  
#from tkinter.ttk import *
from tkinter import ttk
import speech_recognition as sr
from docx import Document
import datetime 
from tkinter import filedialog
from tkinter import messagebox
from tkinter import Canvas
import tkinter.font as font
import sys
from PIL import Image,ImageTk
from deepcorrect import DeepCorrect
#from punctuator import Punctuator

vt="\t"

def doc_create(a):
    document = Document()

    document.add_heading('Speech Recognition', 0)

    p = document.add_paragraph(a)
    date = datetime.datetime.today()
    filename = "ST_"+str(date.day)+str(date.month)+"_"+str(date.hour)+str(date.minute)+".docx"

    document.save(filename)
    txt = "\nYour file is ready...Kindly check "+filename
    messagebox.showinfo("Speech to Text",txt)
    text_field.focus()
    text_field.delete('1.0', END)
    text_field.insert(END, 'Click on Recorded Audio to convert text from audio file...\n')
    text_field.insert(END, 'Click on Start to convert text from microphone...\n')

    #text_field.insert(END, txt)
    #text_field.insert(END, "\nYou can convert other files now...")
    
def audio():
    file =  filedialog.askopenfilename(initialdir = "c:/",title = "Select file",filetypes = (("audio files","*.wav"),("all files","*.*")))
    if file!='': 
        messagebox.showinfo("Speech to Text","Wait till your file is converted")
    
        r = sr.Recognizer()

        harvard = sr.AudioFile(file)
        with harvard as source:
            audio = r.record(source)

        a=r.recognize_google(audio)
        doc_create(a)
    else:
        messagebox.showinfo("Speech to Text","Please select a file")
    
    
def micro():
    global vt
    messagebox.showinfo("Speech to Text","Say Something")
    try:
        r = sr.Recognizer()

        mic = sr.Microphone()

        with mic as source:
            #r.adjust_for_ambient_noise(source)
            audio = r.listen(source)

        b=r.recognize_google(audio)
        #p=Punctuator('c:\python37-32\VoiceToText\punctuator\Demo-Europarl-EN.pcl')
        #print(p.punctuate(b))
        cap = lambda x: x[0].upper() + x[1:]
        cap(b)
        vt = vt + b + ".\n\t"
        
        
        text_field.focus()
        text_field.delete('1.0', END)
        text_field.insert('1.0', vt)
        text_field.insert(END, "\n\nClick on finish if you want to generate a doc file.")
        text_field.insert(END, "\nClick on start if you want to continue.")
        
    except :
        #print("Oops!",sys.exc_info()[0],"occured.")
        text_field.insert(END, "Try Again")

def finished():
    global vt
    if vt!="\t":
        doc_create(vt)
        vt = "\t"
    
def check():
    print("hello")

def show_help():
    print("hello")

root=Tk()  
#App Title  
root.title("Speech to Text Application")

root.title("Speech to Text Application")  
root.attributes("-fullscreen", True)
root.bind("<F11>", lambda event: root.attributes("-fullscreen",
                                    not root.attributes("-fullscreen")))
root.bind("<Escape>", lambda event: root.attributes("-fullscreen", False))
root['bg'] = '#49A'
style=ttk.Style()
style.configure('btn.TButton',borderwidth = 4,relief='solid')
myFont = font.Font(weight="bold",size=20)

ttk.Label(root, text="Select appropriate option", font=(16),background='#49A').pack()
#ttk.Label['bg']='#49A'
#Create Panedwindow  
panedwindow=ttk.Panedwindow(root, orient=HORIZONTAL)  
panedwindow.pack(fill=BOTH, expand=True)  
#Create Frams  
fram1=ttk.Frame(panedwindow,width=300,height=400, relief=SUNKEN)  
fram2=ttk.Frame(panedwindow,width=600,height=400, relief=SUNKEN)  
panedwindow.add(fram1, weight=1)  
panedwindow.add(fram2, weight=4)

im = Image.open("audio_pic.jpeg")
resized = im.resize((100, 100),Image.ANTIALIAS)
tkimage = ImageTk.PhotoImage(resized)
myvar=Label(fram1,image = tkimage)
myvar.image = tkimage
myvar.grid(row=0,column=0,padx=(5,5),pady=(5,5)) 
button1 = ttk.Button(fram1,text="Recorded Audio",command=audio).grid(row =2, column=0, pady =(5,5))

im1 = Image.open("mic_pic1.png")
resized1 = im1.resize((100, 100),Image.ANTIALIAS)
tkimage1 = ImageTk.PhotoImage(resized1)
myvar1=Label(fram1,image = tkimage1)
myvar1.image = tkimage1
myvar1.grid(row=4,column=0,pady=(5,5)) 

#button2 = Button(fram1,text="Microphone",command=start).grid(row =6, column=0, pady =(5,5))
button3 = ttk.Button(fram1,text="Start",command=micro).grid(row = 6, column =0, pady =(5,5))
button4 = ttk.Button(fram1,text="Finish",command=finished).grid(row = 8, column =0, pady =(5,5))
button5 = ttk.Button(fram1,text="Convert",command=check).grid(row = 10, column =0, pady =(5,5))
button6 = ttk.Button(fram1,text="Help",command=show_help).grid(row = 12, column =0, pady =(5,5))
#button3.configure(borderwidth = 2)
#label1 = Label(fram1, text = "Microphone").grid(row=2,)
#button2 = Button(fram1,text="Start",command=micro).grid(row = 4, column =0)
#button2 = Button(fram1,text="Finish",command=finished).grid(row = 4, column =1)
text_field = Text(fram2, height=40, width=60)
text_field.insert(END, 'Click on Recorded Audio to convert text from audio file...\n')
text_field.insert(END, 'Click on Start to convert text from microphone...\n')
text_field.pack(side=LEFT)


#Calling Main()  
root.mainloop() 
