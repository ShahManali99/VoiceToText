#Create app using with separating widget  
from tkinter import *  
from tkinter.ttk import *
from tkinter import ttk
import speech_recognition as sr
from docx import Document
import datetime 
from tkinter import filedialog
from tkinter import messagebox


def audio():
    file =  filedialog.askopenfilename(initialdir = "c:/",title = "Select file",filetypes = (("audio files","*.wav"),("all files","*.*")))
    print(file)
    text1.insert(END, 'Wait till your audio file is converted...\n')
    text1.pack()
    #l1 = Label(fram2, text="Wait while audio file is converted...")
    #l1.place(relx = 0.5, rely = 0.5, anchor = CENTER)
    
    audio_final(file)
    
def audio_final(file):
    r = sr.Recognizer()

    harvard = sr.AudioFile(file)
    with harvard as source:
        audio = r.record(source)

    a=r.recognize_google(audio)

    document = Document()

    document.add_heading('Speech Recognition', 0)

    p = document.add_paragraph(a)
    date = datetime.datetime.today()
    filename = "STA"+str(date.day)+str(date.month)+str(date.hour)+str(date.minute)+".docx"

    document.save(filename)
    txt = "Your file is ready...Kindly check"+filename
    messagebox.showinfo("Speech to Text",txt)
    text1.insert(END, txt)
    text1.insert(END, "\nYou can convert other files now...")
    #l2 = Label(fram2, text=txt)
    #l2.place(relx = 0.5, rely = 0.5, anchor = CENTER)

def micro():
    text1.insert(END, 'Wait till your audio file is converted...\n')
    #Label(fram2, text="Say something").pack()
    micro_final()
    
def micro_final():
    try:
        r = sr.Recognizer()

        mic = sr.Microphone()

        with mic as source:
            audio = r.listen(source)
            #print("Time Over")

        b=r.recognize_google(audio)
        document = Document()

        document.add_heading('Speech Recognition', 0)

        p = document.add_paragraph(b)
        #print(b)
        date = datetime.datetime.today()
        filename = "STM"+str(date.day)+str(date.month)+str(date.hour)+str(date.minute)+".docx"
    
        document.save(filename)
        txt = "Your file is ready...Kindly check"+filename
        messagebox.showinfo("Speech to Text",txt)
        text1.insert(END, txt)
        text1.insert(END, "\nYou can convert other files now...")
        #l3 = Label(fram2, text=txt)
        #l3.place(relx = 0.5, rely = 0.5, anchor = CENTER)
    except:
        text1.insert(END, "Try Again")


root=Tk()  
#App Title  
root.title("Speech to Text Application")

root.title("Speech to Text Application")  
root.attributes("-fullscreen", True)
root.bind("<F11>", lambda event: root.attributes("-fullscreen",
                                    not root.attributes("-fullscreen")))
root.bind("<Escape>", lambda event: root.attributes("-fullscreen", False))

ttk.Label(root, text="Select appropriate option").pack()  
#Create Panedwindow  
panedwindow=ttk.Panedwindow(root, orient=HORIZONTAL)  
panedwindow.pack(fill=BOTH, expand=True)  
#Create Frams  
fram1=ttk.Frame(panedwindow,width=300,height=400, relief=SUNKEN)  
fram2=ttk.Frame(panedwindow,width=600,height=400, relief=SUNKEN)  
panedwindow.add(fram1, weight=1)  
panedwindow.add(fram2, weight=4)

button1 = Button(fram1,text="Recorded Audio",command=audio).grid(row =0, column=0, pady =(5,5))
button2 = Button(fram1,text="Microphone",command=micro).grid(row = 2, column =0)
text1 = Text(fram2, height=40, width=60)
text1.insert(END, 'Wait till your file is converted...\n')
text1.pack(side=LEFT)


#Calling Main()  
root.mainloop() 
